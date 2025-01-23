from concurrent.futures import ThreadPoolExecutor
import os
import json
import requests
import numpy as np
import streamlit as st
from datetime import datetime
from collections import defaultdict
import hashlib
from langchain.text_splitter import RecursiveCharacterTextSplitter
import faiss
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
import fitz  # PyMuPDF
import struct
from docx import Document
from functools import lru_cache
import time  

# Configuración inicial
embedder = SentenceTransformer('all-MiniLM-L6-v2')  # Modelo de embeddings


# Inicialización del estado de sesión
class SessionState:
    def __init__(self):
        dimension = 384  # Dimensión de all-MiniLM-L6-v2
        index = faiss.IndexFlatL2(dimension)  # Índice base (L2 distance)
        self.faiss_index = faiss.IndexIDMap(index)  # Índice que admite IDs personalizados
        self.metadata_map = {}
        self.document_store = defaultdict(list)
        self.chat_history = []
        self.uploaded_files = []

def init_session():
    if 'state' not in st.session_state:
        st.session_state.state = SessionState()

# Función para generar un ID único de 64 bits
def generate_doc_id(file_name, chunk_index):
    hash_object = hashlib.sha256(f"{file_name}_{chunk_index}".encode())
    hash_bytes = hash_object.digest()
    hash_bytes_truncated = hash_bytes[:8]
    doc_id = struct.unpack('>q', hash_bytes_truncated)[0]
    return doc_id

# Clase DocumentProcessor
class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", ". ", "? ", "! ", " "]
        )
    
    def process_file(self, file):
        try:
            # Simular un procesamiento lento (esto es solo para pruebas)
            time.sleep(2)  # Simula un procesamiento de 2 segundos por archivo
            if file is None or file.size == 0:
               st.error(f"El archivo '{file.name}' está vacío o no se cargó correctamente.")
            return False
            # Extraer metadatos del archivo
            if file.name.endswith(".pdf"):
                metadata = extract_pdf_metadata(file)
            elif file.name.endswith(".docx"):
                metadata = extract_docx_metadata(file)
            else:
                st.warning(f"Formato no soportado: {file.name}")
                return False
            
            # Extraer texto del archivo
            text = self._extract_text(file)
            
            # Dividir el texto en chunks
            chunks = self.text_splitter.split_text(text)
            
            # Generar embeddings para cada chunk (en paralelo)
            embeddings = self._get_embeddings(chunks)
            
            # Listas para almacenar embeddings y IDs
            embeddings_list = []
            ids_list = []
            
            # Almacenar metadatos y chunks
            for idx, (chunk, embed) in enumerate(zip(chunks, embeddings)):
                # Generar un ID único de 64 bits
                doc_id = generate_doc_id(file.name, idx)
                
                # Crear metadatos
                chunk_metadata = {
                    **metadata,  # Incluir metadatos del documento
                    "doc_id": doc_id,
                    "content": chunk,
                    "embedding": embed,
                    "source": file.name,
                    "timestamp": datetime.now().isoformat(),
                    "semantic_tags": self._generate_semantic_tags(chunk)
                }
                
                # Almacenar en el estado de la aplicación
                st.session_state.state.metadata_map[doc_id] = chunk_metadata
                st.session_state.state.document_store[file.name].append(doc_id)
                
                # Agregar embeddings e IDs a las listas
                embeddings_list.append(embed)
                ids_list.append(doc_id)
            
            # Actualizar el índice FAISS una sola vez por documento
            embeddings_array = np.array(embeddings_list, dtype=np.float32)
            ids_array = np.array(ids_list, dtype=np.int64)
            st.session_state.state.faiss_index.add_with_ids(embeddings_array, ids_array)
            
            # Mostrar metadatos en la interfaz
            st.success(f"Documento procesado: {metadata['title']}")
            st.write(f"Autor: {metadata['author']}")
            st.write(f"Fecha de creación: {metadata['creation_date']}")
            
            return True
        except Exception as e:
            st.error(f"Error procesando archivo: {str(e)}")
            return False
    
    def _extract_text(self, file):
        if file.name.endswith(".pdf"):
            return self._extract_text_from_pdf(file)
        elif file.name.endswith(".docx"):
            return self._extract_text_from_docx(file)
        else:
            return ""
    
    def _extract_text_from_pdf(self, file):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text("text")  # Modo rápido sin análisis de layout
        return text    
    
    def _extract_text_from_docx(self, file):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _get_embeddings(self, chunks):
        with ThreadPoolExecutor() as executor:
            embeddings = list(executor.map(
                lambda chunk: embedder.encode(chunk, convert_to_tensor=False),
                chunks
            ))
        return embeddings

    @lru_cache(maxsize=1000)
    def _generate_semantic_tags(self, text):
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "deepseek-chat",
                "messages": [{
                    "role": "user",
                    "content": f"Genera 3-5 etiquetas temáticas clave para este texto: {text}"
                }]
            }
        )
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"].split(", ")
        else:
            return []

# Funciones para extraer metadatos
def extract_pdf_metadata(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    metadata = doc.metadata
    return {
        "title": metadata.get("title", "Sin título"),
        "author": metadata.get("author", "Desconocido"),
        "creation_date": metadata.get("creationDate", "N/A")
    }

def extract_docx_metadata(file):
    doc = Document(file)
    core_properties = doc.core_properties
    return {
        "title": core_properties.title or "Sin título",
        "author": core_properties.author or "Desconocido",
        "creation_date": core_properties.created.strftime("%Y-%m-%d") if core_properties.created else "N/A"
    }

# Clase ChatManager
class ChatManager:
    def __init__(self):
        self.vectorizer = TfidfVectorizer(stop_words='english')
    
    def hybrid_search(self, query, top_k=5):
        query_embed = embedder.encode([query])[0]  # Embedding de la consulta
        
        D, I = st.session_state.state.faiss_index.search(
            np.array([query_embed], dtype=np.float32), 
            top_k * 2
        )
        
        lexical_ids = []
        corpus = [" ".join(doc['content']) for doc in st.session_state.state.metadata_map.values()]
        if corpus:
            self.vectorizer.fit(corpus)
            query_vec = self.vectorizer.transform([query])
            doc_scores = cosine_similarity(query_vec, self.vectorizer.transform(corpus)).flatten()
            lexical_ids = np.argsort(doc_scores)[-top_k * 2:][::-1]
        
        combined_ids = list(set(I.flatten().tolist() + lexical_ids))
        valid_ids = [
            idx for idx in combined_ids 
            if 0 <= idx < len(st.session_state.state.metadata_map)
        ]
        
        return [
            st.session_state.state.metadata_map[list(st.session_state.state.metadata_map.keys())[idx]]
            for idx in valid_ids[:top_k]
        ]
    
    def generate_response(self, query, context, sources):
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "deepseek-chat",
                "messages": [{
                    "role": "system",
                    "content": f"Contexto:\n{context}\n\nResponde de manera precisa y profesional."
                }, {
                    "role": "user",
                    "content": query
                }]
            }
        )
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"], sources
        else:
            raise Exception(f"Error en la API: {response.status_code} - {response.text}")


# Interfaz de usuario mejorada
class DeepSeekUI:
    def __init__(self):
        self.processor = DocumentProcessor()
        self.chat_manager = ChatManager()
    
    def render_sidebar(self):
        with st.sidebar:
            st.title("⚙️ Configuración")
            
            # Botón para cambiar entre modo claro y oscuro
            dark_mode = st.checkbox("Modo Oscuro")
            
            if dark_mode:
                st.markdown("""
                <style>
                    .main { background-color: #1e1e1e; color: white; }
                    .stButton>button { background-color: #4CAF50; color: white; }
                    .response-box { 
                        border: 2px solid #4CAF50;
                        border-radius: 5px;
                        padding: 20px;
                        margin: 10px 0;
                        background-color: #2e2e2e;
                    }
                    .reference-item { 
                        margin: 5px 0;
                        padding: 10px;
                        background-color: #3e3e3e;
                        border-radius: 3px;
                    }
                </style>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <style>
                    .main { background-color: #f0f2f6; color: black; }
                    .stButton>button { background-color: #4CAF50; color: white; }
                    .response-box { 
                        border: 2px solid #4CAF50;
                        border-radius: 5px;
                        padding: 20px;
                        margin: 10px 0;
                        background-color: #ffffff;
                    }
                    .reference-item { 
                        margin: 5px 0;
                        padding: 10px;
                        background-color: #e8f5e9;
                        border-radius: 3px;
                    }
                </style>
                """, unsafe_allow_html=True)
            
            with st.expander("📤 Gestión de Documentos"):
                uploaded_files = st.file_uploader(
                    "Subir documentos (PDF/DOCX)",
                    type=["pdf", "docx"],
                    accept_multiple_files=True
                )
                if st.button("Procesar Documentos", type="primary"):
                    if uploaded_files:
                        progress_bar = st.progress(0)  # Barra de progreso
                        status_text = st.empty()  # Texto de estado
                        
                        for i, file in enumerate(uploaded_files):
                            status_text.text(f"Procesando {i + 1}/{len(uploaded_files)}: {file.name}")
                            if file.name not in st.session_state.state.uploaded_files:
                                try:
                                    if not self.processor.process_file(file):
                                        st.warning(f"No se pudo procesar el archivo: {file.name}")
                                except Exception as e:
                                    st.error(f"Error procesando archivo '{file.name}': {str(e)}")
                            progress_bar.progress((i + 1) / len(uploaded_files))  # Actualizar barra de progreso  
                                                  
                        status_text.text("Procesamiento completado.")
                        st.success("Todos los documentos han sido procesados.")
                    else:
                        st.warning("No se han subido archivos.")
                                    
            with st.expander("🔍 Opciones de Búsqueda"):
                self.search_type = st.radio(
                    "Modo de búsqueda:",
                    ["Híbrida", "Vectorial", "Semántica"],
                    index=0
                )
                
                self.creativity = st.slider(
                    "Nivel de creatividad:",
                    min_value=0.0, max_value=1.0, value=0.5,
                    help="Controla el balance entre precisión y originalidad"
                )
    
    def render_chat(self):
        st.title("🧠 Asistente DeepSeek RAG")
        
        # Historial de chat
        chat_container = st.container(height=500)
        with chat_container:
            for msg in st.session_state.state.chat_history:
                self._render_message(msg)
        
        # Entrada de usuario
        query = st.chat_input("Escribe tu pregunta...")
        if query:
            self._handle_user_query(query)
    
    def _render_message(self, msg):
        if msg['type'] == 'user':
            st.markdown(f"""
            <div style="margin: 1rem 0; padding: 1rem; 
                        background: #e3f2fd; border-radius: 10px;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1)">
                <strong>👤 Tú:</strong> {msg['content']}
            </div>
            """, unsafe_allow_html=True)
        
        elif msg['type'] == 'assistant':
            with st.expander("💡 Respuesta Completa", expanded=True):
                st.markdown(f"""
                <div style="margin: 0.5rem 0; padding: 1rem;
                            background: #f5f5f5; border-radius: 10px;">
                    <strong>🤖 Asistente:</strong> {msg['content']}
                </div>
                """, unsafe_allow_html=True)
                
                if 'references' in msg:
                    st.markdown("**🔍 Fuentes Relacionadas:**")
                    for ref in msg['references']:
                        st.markdown(f"""
                        <div class="reference-item">
                            📌 {ref}
                        </div>
                        """, unsafe_allow_html=True)
    
    def _handle_user_query(self, query):
        st.session_state.state.chat_history.append({
            'type': 'user',
            'content': query
        })
        
        with st.spinner("🔍 Buscando información relevante..."):
            results = self.chat_manager.hybrid_search(query)
            context = "\n\n".join(
                f"[Fuente: {res['source']}]\n{res['content']}" 
                for res in results
            )
            sources = [f"{res['source']} - {', '.join(res['semantic_tags'][:2])}" for res in results]
        
        with st.spinner("💡 Generando respuesta..."):
            response, response_sources = self.chat_manager.generate_response(query, context, sources)
            
            st.session_state.state.chat_history.append({
                'type': 'assistant',
                'content': response,
                'references': response_sources,  # Incluye las fuentes en la respuesta
                'validation': self._validate_response(response, context)
            })
        
        st.rerun()    
    def _validate_response(self, response, context):
        validation_prompt = f"""
        Evalúa la siguiente respuesta basada en el contexto proporcionado:
        
        **Respuesta:**
        {response}
        
        **Contexto:**
        {context}
        
        Proporciona una validación en formato JSON con:
        - score (1-5)
        - valid (true/false)
        - reasons (lista de razones)
        """
        
        validation = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "deepseek-chat",
                "messages": [{
                    "role": "user",
                    "content": validation_prompt
                }],
                "response_format": {"type": "json_object"}
            }
        )
        if validation.status_code == 200:
            return json.loads(validation.json()["choices"][0]["message"]["content"])
        else:
            return {"score": 0, "valid": False, "reasons": ["Error en la validación"]}

# Inicialización y ejecución
if __name__ == "__main__":
    init_session()
    ui = DeepSeekUI()
    ui.render_sidebar()
    ui.render_chat()