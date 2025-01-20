import fitz
import streamlit as st
import io
import os
import tempfile
import pyperclip
from docx import Document

def pdf_to_text_with_page_numbers(uploaded_file):
    try:
        file_bytes = uploaded_file.read()
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_filename = temp_file.name
            temp_file.write(file_bytes)
            temp_file.flush()

        try:
            doc = fitz.open(temp_filename)
            metadata = doc.metadata
            title = metadata.get("title", "Sin título")
            author = metadata.get("author", "Desconocido")
            text_with_page_numbers = []
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text("text", sort=True)
                if text.strip():
                    text_with_page_numbers.append(f"Página {page_num + 1}:\n{text}\n")
                else:
                    st.info(f"Página {page_num + 1}: No se encontró texto en esta página.")
            return title, author, text_with_page_numbers
        finally:
            try:
                os.remove(temp_filename)
            except PermissionError:
                st.error("Error al eliminar el archivo temporal. Puede que esté siendo utilizado por otro proceso.")

    except fitz.FileDataError as e:
        st.error(f"Error: El archivo no es un PDF válido o está corrupto: {e}")
        return None, None, None
    except Exception as e:
        st.error(f"Error general al procesar el PDF: {e}")
        return None, None, None

def save_text_to_docx(title, author, text_with_pages, output_filename):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_heading(f"Autor: {author}", level=2)
    for page_text in text_with_pages:
        doc.add_paragraph(page_text)
    doc.save(output_filename)

def main():
    st.title("Depurador de Libros PDF para RAG")
    st.subheader("Extrae texto con números de página y guarda en DOCX")

    uploaded_file = st.file_uploader("Cargar archivo PDF", type="pdf")

    if uploaded_file is not None:
        title, author, text_with_pages = pdf_to_text_with_page_numbers(uploaded_file)

        if text_with_pages:
            st.header("Texto extraído del PDF:")
            all_text = ""
            for page_text in text_with_pages:
                st.write(page_text)
                st.write("-" * 80)
                all_text += page_text

            output_filename = "output.docx"
            save_text_to_docx(title, author, text_with_pages, output_filename)
            st.success(f"Texto guardado en {output_filename}")

            if st.button("Copiar todo el texto al portapapeles"):
                try:
                    pyperclip.copy(all_text)
                    st.success("Texto copiado al portapapeles!")
                except pyperclip.PyperclipException:
                    st.error("Error al copiar al portapapeles. Asegúrate de tener pyperclip instalado. (pip install pyperclip)")

        elif text_with_pages is None:
            st.warning("No se pudo extraer texto del PDF. Revise el formato del archivo.")

if __name__ == "__main__":
    main()
