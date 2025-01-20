import os
import io
import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from googletrans import Translator
from mistralai.client import MistralClient
from mistralai.exceptions import MistralAPIException
from mistralai.models.chat_completion import ChatMessage
import faiss
import numpy as np
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph
import pyperclip

# Configuración de Mistral
mistral_api_key = "yxrklpJOVs2mhBMfi0UK74CfGpzGcbsI"
mistral = MistralClient(api_key=mistral_api_key)
embed_model = "mistral-embed"

# Configuración de FAISS
dimension = 1024
index = faiss.IndexFlatL2(dimension)  # Usamos L2 como métrica de distancia

# Traductor
translator = Translator()

def translate_to_english(text):
    try:
        return translator.translate(text, src='es', dest='en').text
    except Exception as e:
        st.error(f"Error al traducir al inglés: {e}")
        return text

def translate_to_spanish(text):
    try:
        return translator.translate(text, src='en', dest='es').text
    except Exception as e:
        st.error(f"Error al traducir al español: {e}")
        return text

# Función para leer PDFs y extraer el título y metadatos
def read_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        metadata = doc.metadata
        title = metadata.get("title", "Sin título") or os.path.basename(file_path)
        author = metadata.get("author", "Desconocido")
        year = metadata.get("creationDate", "N/A")[:4]  # Extraer el año de la fecha de creación
        text = "".join([doc.load_page(page_num).get_text() for page_num in range(len(doc))])
        return title, author, year, text, metadata
    except Exception as e:
        st.error(f"Error al leer el archivo PDF: {e}")
        return "Error al leer el archivo", "", "", "", {}

# Función para leer archivos Word y extraer el título y metadatos
def read_docx(file_path):
    try:
        doc = Document(file_path)
        title = doc.core_properties.title or os.path.basename(file_path)
        author = doc.core_properties.author or "Desconocido"
        year = doc.core_properties.created.year
        text = "\n".join([para.text for para in doc.paragraphs])
        metadata = {
            "title": title,
            "author": author,
            "year": year,
            "pages": len(doc.paragraphs)
        }
        return title, author, year, text, metadata
    except Exception as e:
        st.error(f"Error al leer el archivo Word: {e}")
        return "Error al leer el archivo", "", "", "", {}

# Función de embedding
def embed(metadata):
    batch_size = len(metadata)
    while batch_size >= 2:
        try:
            embeds = []
            for j in range(0, len(metadata), batch_size):
                j_end = min(len(metadata), j + batch_size)
                embed_response = mistral.embeddings(
                    input=[x["title"] + "\n" + x["content"] for x in metadata[j:j_end]],
                    model=embed_model
                )
                embeds.extend([x.embedding for x in embed_response.data])
            return embeds
        except MistralAPIException as e:
            st.error(f"Error al generar embeddings: {e}")
            batch_size //= 2
    raise MistralAPIException("Error al generar embeddings")

# Función de recuperación con metadatos
def get_docs_with_metadata(query, top_k):
    query_en = translate_to_english(query)
    xq = mistral.embeddings(input=[query_en], model=embed_model).data[0].embedding
    D, I = index.search(np.array([xq], dtype=np.float32), top_k)

    docs = []
    for i in I[0]:
        if 0 <= i < len(data):
            doc_content = data[i]["metadata"]["content"]
            doc_metadata = data[i]["metadata"]
            docs.append((doc_content, doc_metadata))  # Incluye los metadatos junto al contenido

    return docs

# Función de generación con referencias
def generate_with_references(query, docs):
    query_en = translate_to_english(query)
    system_message = (
        "te encargaras de darle forma ala repuesta que han solicitado, para esto se daran pedazos de informacion. "
        "esos embeding te permitiran desarrollar una respuesta coherente y precisa. "
        "de ahi obtendras el autor de las fuentes que semanticamente son mas proximas a la pregunta.\n\n"
        "CONTEXT:\n" + "\n---\n".join([doc[0] for doc in docs])  # Solo el contenido del documento
    )
    messages = [
        ChatMessage(role="system", content=system_message),
        ChatMessage(role="user", content=query_en)
    ]
    chat_response = mistral.chat(
        model="mistral-large-latest",
        messages=messages
    )
    response_en = chat_response.choices[0].message.content

    # Construir las citas para las fuentes
    references = []
    for doc_content, doc_metadata in docs:
        reference = f"Fuente: {doc_metadata['title']} ({doc_metadata['year']}), Autor: {doc_metadata['author']}, Página(s): {doc_metadata.get('pages', 'N/A')}"
        references.append(reference)

    # Unir la respuesta con las referencias
    response_with_references = f"{response_en}\n\nReferencias:\n" + "\n".join(references)

    return translate_to_spanish(response_with_references)

# Función para generar el PDF con referencias
def generate_response_pdf_with_references(response):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
    )
    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(name='Normal', fontName='Helvetica', fontSize=12, leading=14)
    story = [Paragraph(response, normal_style)]
    doc.build(story)
    buffer.seek(0)
    return buffer

# Función para copiar al portapapeles
def copy_to_clipboard(text):
    try:
        pyperclip.copy(text)
        st.success("Texto copiado al portapapeles")
    except pyperclip.PyperclipException as e:
        st.error(f"No se pudo copiar al portapapeles: {e}")

# Interfaz de Usuario con Streamlit
st.title("Sistema de Recuperación Aumentada de Generación (RAG)")

# Variable de sesión para almacenar los títulos de los archivos cargados
if 'uploaded_file_titles' not in st.session_state:
    st.session_state.uploaded_file_titles = []

if 'uploaded_file_contents' not in st.session_state:
    st.session_state.uploaded_file_contents = {}

uploaded_files = st.file_uploader("Cargar archivos PDF o Word", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    for uploaded_file in uploaded_files:
        file_path = os.path.join(os.getcwd(), uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        if uploaded_file.name.endswith(".pdf"):
            title, author, year, text, metadata = read_pdf(file_path)
        elif uploaded_file.name.endswith(".docx"):
            title, author, year, text, metadata = read_docx(file_path)
        else:
            st.error("Formato de archivo no soportado")
            title, author, year, text, metadata = "Formato no soportado", "", "", "", {}

        if title:
            st.write(f"Título del archivo cargado: {title}")
            st.write(f"Autor: {author}")
            st.write(f"Año: {year}")
            st.session_state.uploaded_file_titles.append(title)
            st.session_state.uploaded_file_contents[title] = text

            # Procesar el texto extraído
            data = [{"id": str(len(st.session_state.uploaded_file_titles)), "metadata": {"title": title, "content": text, "year": year, "author": author, "pages": metadata.get('pages', 'N/A')}}]
            try:
                with st.spinner('Procesando...'):
                    embeds = embed(data)
                    index.add(np.array(embeds, dtype=np.float32))
            except MistralAPIException as e:
                st.error(f"Error al generar embeddings: {e}")

query = st.text_input("Introduce tu consulta en español:")
if st.button("Obtener Respuesta"):
    if query:
        with st.spinner('Buscando respuesta...'):
            docs_with_metadata = get_docs_with_metadata(query, top_k=5)  # Obtener documentos con metadatos
            response_with_references = generate_with_references(query=query, docs=docs_with_metadata)
            st.write("Respuesta:", response_with_references)

            # Botones para descargar o copiar la respuesta
            col1, col2 = st.columns(2)
            with col1:
                pdf_buffer = generate_response_pdf_with_references(response_with_references)
                st.download_button(
                    label="Descargar respuesta con referencias en PDF",
                    data=pdf_buffer,
                    file_name="respuesta_con_referencias.pdf",
                    mime="application/pdf"
                )
            with col2:
                st.text_area("Texto de la respuesta:", response_with_references, height=200)
                st.info("Selecciona el texto y cópialo manualmente si el botón no funciona.")

    else:
        st.warning("Por favor, introduce una consulta.")
