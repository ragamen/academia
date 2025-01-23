from concurrent.futures import ThreadPoolExecutor
import os
import json
import requests
import numpy as np
import streamlit as st
from datetime import datetime
from collections import defaultdict
import hashlib
import re
import fitz  # PyMuPDF
import struct
from docx import Document
# Agrega estas importaciones al inicio del archivo
from docx.table import Table
from docx.document import Document as DocxDocument
from io import BytesIO
from functools import lru_cache
import time
from langchain.text_splitter import RecursiveCharacterTextSplitter
import faiss
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from docx.text.paragraph import Paragraph  # Importar Paragraph

# Configuración inicial
embedder = SentenceTransformer('all-MiniLM-L6-v2')


# Configurar el tema oscuro
st.set_page_config(page_title="Plataforma RAG", page_icon="🤖", layout="centered", initial_sidebar_state="auto")

# Aplicar estilo CSS para el tema oscuro
st.markdown(
    """
    <style>
    .main { background-color: #1e1e1e; color: white; }
    .stButton>button { background-color: #4CAF50; color: white; }
    .response-box { 
        border: 2px solid #4CAF50;
        border-radius: 5px;
        padding: 20px;
        margin: 10px 0;
        background-color: #2e2e2e;
    }
    .reference-item { 
        margin: 5px 0;
        padding: 10px;
        background-color: #3e3e3e;
        border-radius: 3px;
    }
    </style>
    """,
    unsafe_allow_html=True
)
class SessionState:
    def __init__(self):
        dimension = 384
        index = faiss.IndexFlatL2(dimension)
        self.faiss_index = faiss.IndexIDMap(index)
        self.metadata_map = {}
        self.document_store = defaultdict(list)
        self.chat_history = []
        self.uploaded_files = []
        self.current_page = None  # Nueva variable para seguimiento de página

def init_session():
    if 'state' not in st.session_state:
        st.session_state.state = SessionState()

def generate_doc_id(file_name, chunk_index):
    hash_object = hashlib.sha256(f"{file_name}_{chunk_index}".encode())
    return struct.unpack('>q', hash_object.digest()[:8])[0]

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", ". ", "? ", "! ", " "]
        )
    
    def process_file(self, file):
        try:
            if not file or file.size == 0:
                raise ValueError("Archivo vacío o inválido")
            
            file_content = file.read()
            file.seek(0)
            
            if file.name.endswith(".pdf"):
                metadata = self._extract_pdf_metadata(file_content)
                text = self._extract_pdf_text(file_content)
            elif file.name.endswith(".docx"):
                metadata = self._extract_docx_metadata(file)
                text = self._extract_docx_text(file_content)
            else:
                raise ValueError("Formato no soportado")
            
            chunks = self.text_splitter.split_text(text)
            embeddings = self._generate_embeddings(chunks)
            
            self._store_chunks(file.name, chunks, embeddings, metadata)
            
            return True
            
        except Exception as e:
            st.error(f"Error procesando {file.name}: {str(e)}")
            return False

    def _extract_pdf_metadata(self, file_content):
        try:
            with fitz.open(stream=file_content, filetype="pdf") as doc:
                return {
                    "title": doc.metadata.get("title", "Sin título"),
                    "author": doc.metadata.get("author", "Desconocido"),
                    "creation_date": doc.metadata.get("creationDate", "N/A"),
                    "pages": len(doc)
                }
        except Exception as e:
            st.error(f"Error PDF: {str(e)}")
            return {"title": "Desconocido", "author": "Desconocido", "creation_date": "N/A", "pages": 0}

    def _extract_pdf_text(self, file_content):
        try:
            text = []
            with fitz.open(stream=file_content, filetype="pdf") as doc:
                for page_num in range(len(doc)):
                    text.append(f"\x02PAGE:{page_num+1}\x03{doc.load_page(page_num).get_text('text')}")
            return "\n".join(text)
        except Exception as e:
            st.error(f"Error leyendo PDF: {str(e)}")
            return ""

    def _extract_docx_metadata(self, file):
        """Extrae metadatos de DOCX sin páginas."""
        try:
            doc = Document(file)
            props = doc.core_properties
            
            return {
                "titulo": props.title or file.name,
                "autor": props.author or "Desconocido",
                "fecha_creacion": str(props.created) if props.created else "N/A"
            }
        except Exception as e:
            st.error(f"Error DOCX: {str(e)}")
            return {
                "titulo": file.name,
                "autor": "Desconocido",
                "fecha_creacion": "N/A"
            }
            
    def _extract_docx_text(self, file_content):
        try:
            doc = Document(BytesIO(file_content))
            text = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            return "\n".join(text)
        except Exception as e:
            st.error(f"Error leyendo DOCX: {str(e)}")
            return ""

    def _generate_embeddings(self, chunks):
        with ThreadPoolExecutor() as executor:
            return list(executor.map(
                lambda chunk: embedder.encode(chunk, convert_to_tensor=False),
                chunks
            ))

def _store_chunks(self, filename, chunks, embeddings, metadata):
    ids = []
    for idx, (chunk, embed) in enumerate(zip(chunks, embeddings)):
        doc_id = generate_doc_id(filename, idx)
        
        # Eliminar búsqueda de páginas para DOCX
        if filename.endswith(".docx"):
            cleaned_chunk = chunk  # DOCX no tiene marcadores de página
            page_info = {}  # No incluir campo de página
        else:
            page_match = re.search(r'\x02PAGE:(\d+)\x03', chunk)
            cleaned_chunk = re.sub(r'\x02PAGE:\d+\x03', '', chunk)
            page_info = {"page": page_match.group(1)} if page_match else {}
        
        st.session_state.state.metadata_map[doc_id] = {
            "doc_id": doc_id,
            "content": cleaned_chunk,
            "embedding": embed,
            "source": filename,
            "metadata": {
                **metadata,
                **page_info  # Solo PDF tendrá este campo
            }
        }
        ids.append(doc_id)
    
    embeddings = np.array([e for e in embeddings if e is not None], dtype=np.float32)
    ids = np.array(ids, dtype=np.int64)
    st.session_state.state.faiss_index.add_with_ids(embeddings, ids)
    st.session_state.state.document_store[filename].extend(ids)

            
            
            
# Clase ChatManager
class ChatManager:
    def __init__(self):
        self.vectorizer = TfidfVectorizer(stop_words='english')
    
    def hybrid_search(self, query, top_k=3):
        try:
            # Búsqueda vectorial estricta
            query_embed = embedder.encode([query])[0]
            D, I = st.session_state.state.faiss_index.search(
                np.array([query_embed], dtype=np.float32), 
                top_k * 2
            )
            
            # Filtrado riguroso
            valid_results = []
            for i, distance in zip(I[0], D[0]):
                if i in st.session_state.state.metadata_map:
                    metadata = st.session_state.state.metadata_map[i]
                    similarity = 1 / (1 + distance)
                    if similarity >= self.similarity_threshold and len(metadata['content']) >= self.min_chunk_length:
                        valid_results.append(metadata)
            
            return valid_results[:top_k]
        
        except Exception as e:
            st.error(f"Error de búsqueda: {str(e)}")
            return []

    def generate_response(self, query):
        results = self.hybrid_search(query)
        
        # Validación definitiva
        if not results:
            return "No tengo información documental sobre este tema.", []
        
        # Prompt de seguridad
        system_prompt = """
        Eres un asistente especializado en documentos internos. 
        Responde EXCLUSIVAMENTE con la información proporcionada. 
        Si el contexto no menciona el tema, di: 'No hay información relevante en los documentos'.
        """
        
        context = "\n\n".join(f"[Fuente: {res['source']}]\n{res['content']}" for res in results)
        
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
            json={
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Contexto:\n{context}\n\nPregunta: {query}"}
                ],
                "temperature": 0.1  # Reducir creatividad
            }
        )
        
        return response.json()["choices"][0]["message"]["content"], []
        
        
# Interfaz de usuario mejorada
class DeepSeekUI:
    def __init__(self):
        self.processor = DocumentProcessor()
        self.chat_manager = ChatManager()
    def _render_metadata(self):
        """Muestra metadatos sin páginas para DOCX."""
        with st.sidebar.expander("📚 Referencias Documentales", expanded=True):
            if not st.session_state.state.uploaded_files:
                st.info("Sube documentos para ver sus metadatos.")
                return
                
            for file_name in st.session_state.state.uploaded_files:
                doc_ids = st.session_state.state.document_store.get(file_name, [])
                if not doc_ids:
                    continue
                    
                chunk_data = st.session_state.state.metadata_map.get(doc_ids[0], {})
                metadata = chunk_data.get("metadata", {})
                
                st.markdown(f"**Documento:** `{file_name}`")
                st.markdown(f"**Título:** {metadata.get('titulo', 'N/A')}")
                st.markdown(f"**Autor:** {metadata.get('autor', 'N/A')}")
                if "page" in metadata:  # Solo para PDF
                    st.markdown(f"**Página:** {metadata['page']}")
                st.markdown("---")
    def render_sidebar(self):
        with st.sidebar:
            st.title("⚙️ Configuración")
            
            # Botón para cambiar entre modo claro y oscuro
            dark_mode = st.checkbox("Modo Oscuro")
            
            if dark_mode:
                st.markdown("""
                <style>
                    .main { background-color: #1e1e1e; color: white; }
                    .stButton>button { background-color: #4CAF50; color: white; }
                    .response-box { 
                        border: 2px solid #4CAF50;
                        border-radius: 5px;
                        padding: 20px;
                        margin: 10px 0;
                        background-color: #2e2e2e;
                    }
                    .reference-item { 
                        margin: 5px 0;
                        padding: 10px;
                        background-color: #3e3e3e;
                        border-radius: 3px;
                    }
                </style>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <style>
                    .main { background-color: #f0f2f6; color: black; }
                    .stButton>button { background-color: #4CAF50; color: white; }
                    .response-box { 
                        border: 2px solid #4CAF50;
                        border-radius: 5px;
                        padding: 20px;
                        margin: 10px 0;
                        background-color: #ffffff;
                    }
                    .reference-item { 
                        margin: 5px 0;
                        padding: 10px;
                        background-color: #e8f5e9;
                        border-radius: 3px;
                    }
                </style>
                """, unsafe_allow_html=True)
            
            with st.expander("📤 Gestión de Documentos"):
                uploaded_files = st.file_uploader(
                    "Subir documentos (PDF/DOCX)",
                    type=["pdf", "docx"],
                    accept_multiple_files=True
                )
                if st.button("Procesar Documentos", type="primary"):
                    if uploaded_files:
                        progress_bar = st.progress(0)  # Barra de progreso
                        status_text = st.empty()  # Texto de estado
                        
                        for i, file in enumerate(uploaded_files):
                            status_text.text(f"Procesando {i + 1}/{len(uploaded_files)}: {file.name}")
                            if file.name not in st.session_state.state.uploaded_files:
                                try:
                                    if not self.processor.process_file(file):
                                        st.warning(f"No se pudo procesar el archivo: {file.name}")
                                except Exception as e:
                                    st.error(f"Error procesando archivo '{file.name}': {str(e)}")
                            progress_bar.progress((i + 1) / len(uploaded_files))  # Actualizar barra de progreso  
                                                  
                        status_text.text("Procesamiento completado.")
                        st.success("Todos los documentos han sido procesados.")
                    else:
                        st.warning("No se han subido archivos.")
                                    
            with st.expander("🔍 Opciones de Búsqueda"):
                self.search_type = st.radio(
                    "Modo de búsqueda:",
                    ["Híbrida", "Vectorial", "Semántica"],
                    index=0
                )
                
                self.creativity = st.slider(
                    "Nivel de creatividad:",
                    min_value=0.0, max_value=1.0, value=0.5,
                    help="Controla el balance entre precisión y originalidad"
                )
    
    def render_chat(self):
        st.title("🧠 Asistente DeepSeek RAG")
        
        # Historial de chat
        chat_container = st.container(height=500)
        with chat_container:
            for msg in st.session_state.state.chat_history:
                self._render_message(msg)
        
        # Entrada de usuario
        query = st.chat_input("Escribe tu pregunta...")
        if query:
            self._handle_user_query(query)
    
    def _render_message(self, msg):
        if msg['type'] == 'user':
            st.markdown(f"""
            <div style="margin: 1rem 0; padding: 1rem; 
                        background: #e3f2fd; border-radius: 10px;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1)">
                <strong>👤 Tú:</strong> {msg['content']}
            </div>
            """, unsafe_allow_html=True)
        
        elif msg['type'] == 'assistant':
            with st.expander("💡 Respuesta Completa", expanded=True):
                st.markdown(f"""
                <div style="margin: 0.5rem 0; padding: 1rem;
                            background: #f5f5f5; border-radius: 10px;">
                    <strong>🤖 Asistente:</strong> {msg['content']}
                </div>
                """, unsafe_allow_html=True)
                
                if 'references' in msg:
                    st.markdown("**🔍 Fuentes Relacionadas:**")
                    for ref in msg['references']:
                        st.markdown(f"""
                        <div class="reference-item">
                            📌 {ref}
                        </div>
                        """, unsafe_allow_html=True)
    
    def _handle_user_query(self, query):
        st.session_state.state.chat_history.append({
            'type': 'user',
            'content': query
        })
        
        with st.spinner("🔍 Buscando información relevante..."):
            results = self.chat_manager.hybrid_search(query)
            context = "\n\n".join(
                f"[Fuente: {res['source']}]\n{res['content']}" 
                for res in results
            )
            sources = [f"{res['source']} - {', '.join(res['semantic_tags'][:2])}" for res in results]
        
        with st.spinner("💡 Generando respuesta..."):
            response, response_sources = self.chat_manager.generate_response(query, context, sources)
            
            st.session_state.state.chat_history.append({
                'type': 'assistant',
                'content': response,
                'references': response_sources,  # Incluye las fuentes en la respuesta
                'validation': self._validate_response(response, context)
            })
        
        st.rerun()    
        
    def _validate_response(self, response, context):
        try:
            # Eliminar esta validación externa si no es esencial
            return {"status": "ok"}  # Simular validación exitosa
            
            # Opcional: Si realmente necesitas validación, usa esto:
            validation_prompt = f"Valida si esta respuesta es coherente con el contexto: {response}"
            
            # Añadir timeout y reintentos
            for _ in range(3):  # 3 reintentos
                try:
                    validation = requests.post(
                        "https://api.deepseek.com/v1/chat/completions",
                        headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"},
                        json={
                            "model": "deepseek-chat",
                            "messages": [{"role": "user", "content": validation_prompt}]
                        },
                        timeout=10  # 10 segundos máximo
                    )
                    validation.raise_for_status()  # Lanza error si HTTP != 200
                    return json.loads(validation.json()["choices"][0]["message"]["content"])
                
                except (requests.exceptions.ChunkedEncodingError, 
                        requests.exceptions.Timeout,
                        requests.exceptions.ConnectionError):
                    time.sleep(2)  # Esperar 2 segundos antes de reintentar
                    continue
                    
            return {"error": "Falló la validación después de 3 intentos"}
            
        except Exception as e:
            return {"error": str(e)}
# Inicialización y ejecución
if __name__ == "__main__":
    init_session()
    ui = DeepSeekUI()
    ui.render_sidebar()
    ui.render_chat()